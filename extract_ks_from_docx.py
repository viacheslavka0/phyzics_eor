#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Скрипт для извлечения тем "Система знаний" из docx файла
"""

import re
from docx import Document
import json

def extract_knowledge_systems(docx_path):
    """Извлекает все упоминания систем знаний из docx файла"""
    
    doc = Document(docx_path)
    knowledge_systems = []
    
    # Паттерны для поиска
    patterns = [
        r'Система знаний\s*[«"](.+?)[»"]',  # Система знаний «...»
        r'Система знаний\s*по\s*теме\s*[«"](.+?)[»"]',  # Система знаний по теме «...»
        r'Система знаний\s*по\s*теме\s*(.+?)(?:\.|$|\n)',  # Система знаний по теме ...
        r'Система знаний\s*«(.+?)»',  # Система знаний «...» (без пробела)
    ]
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
            
        # Ищем все паттерны
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                ks_title = match.group(1).strip()
                # Очищаем от кавычек и лишних пробелов
                ks_title = re.sub(r'^[«"]+|[»"]+$', '', ks_title).strip()
                if ks_title:
                    # Проверяем на дубликаты (без учёта кавычек)
                    normalized_title = re.sub(r'[«»""]', '', ks_title).strip()
                    if normalized_title not in [re.sub(r'[«»""]', '', ks['title']).strip() for ks in knowledge_systems]:
                        knowledge_systems.append({
                            'title': ks_title,
                            'source_text': text[:100] + '...' if len(text) > 100 else text
                        })
    
    # Также проверяем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if not text:
                    continue
                    
                for pattern in patterns:
                    matches = re.finditer(pattern, text, re.IGNORECASE)
                    for match in matches:
                        ks_title = match.group(1).strip()
                        # Очищаем от кавычек и лишних пробелов
                        ks_title = re.sub(r'^[«"]+|[»"]+$', '', ks_title).strip()
                        if ks_title:
                            # Проверяем на дубликаты (без учёта кавычек)
                            normalized_title = re.sub(r'[«»""]', '', ks_title).strip()
                            if normalized_title not in [re.sub(r'[«»""]', '', ks['title']).strip() for ks in knowledge_systems]:
                                knowledge_systems.append({
                                    'title': ks_title,
                                    'source_text': text[:100] + '...' if len(text) > 100 else text
                                })
    
    return knowledge_systems

def main():
    docx_path = r"C:\Users\Вячеслав\Desktop\Магистратура\ВКР\Материалы от Лидии Алексеевны\ВАКО\ВАКО\задания.docx"
    
    print("Извлечение систем знаний из docx файла...")
    print(f"Путь к файлу: {docx_path}\n")
    
    try:
        knowledge_systems = extract_knowledge_systems(docx_path)
        
        print(f"Найдено систем знаний: {len(knowledge_systems)}\n")
        print("=" * 80)
        
        for i, ks in enumerate(knowledge_systems, 1):
            print(f"\n{i}. {ks['title']}")
            print(f"   Источник: {ks['source_text']}")
        
        # Сохраняем в JSON для удобства
        output_file = "knowledge_systems.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(knowledge_systems, f, ensure_ascii=False, indent=2)
        
        print("\n" + "=" * 80)
        print(f"\nРезультаты сохранены в файл: {output_file}")
        print("\nФормат для БД:")
        print("-" * 80)
        for i, ks in enumerate(knowledge_systems, 1):
            print(f"{i}. {ks['title']}")
        
    except FileNotFoundError:
        print(f"Ошибка: Файл не найден по пути {docx_path}")
        print("\nПроверьте путь к файлу и попробуйте снова.")
    except Exception as e:
        print(f"Ошибка: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
