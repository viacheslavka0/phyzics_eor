#!/usr/bin/env python
# -*- coding: utf-8 -*-
import sys
import os

try:
    from docx import Document
except ImportError:
    print("Установите python-docx: pip install python-docx")
    sys.exit(1)

# Путь к файлу
file_path = r"C:\Users\Вячеслав\Desktop\Магистратура\ВКР\ОТЧЕТ О НИР\Разработка сценария работы ЭОР\Сценарий ЭОР_с экранами.docx"

if not os.path.exists(file_path):
    print(f"Файл не найден: {file_path}")
    sys.exit(1)

output_file = "scenario_extracted.txt"

try:
    doc = Document(file_path)
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("СОДЕРЖИМОЕ ДОКУМЕНТА\n")
        f.write("=" * 80 + "\n\n")
        
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:
                f.write(f"{i}. {text}\n")
        
        f.write("\n" + "=" * 80 + "\n")
        f.write("ТАБЛИЦЫ\n")
        f.write("=" * 80 + "\n\n")
        
        for table_idx, table in enumerate(doc.tables, 1):
            f.write(f"Таблица {table_idx}:\n")
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    f.write(f"  {row_text}\n")
            f.write("\n")
        
        f.write("=" * 80 + "\n")
        f.write(f"Всего параграфов: {len([p for p in doc.paragraphs if p.text.strip()])}\n")
        f.write(f"Всего таблиц: {len(doc.tables)}\n")
    
    print(f"✅ Текст извлечен и сохранен в файл: {output_file}")
    print(f"📄 Всего параграфов: {len([p for p in doc.paragraphs if p.text.strip()])}")
    print(f"📊 Всего таблиц: {len(doc.tables)}")
    
except Exception as e:
    print(f"Ошибка при чтении файла: {e}")
    import traceback
    traceback.print_exc()
